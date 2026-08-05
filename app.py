import base64
import io
import json
import math
import os
import re
from datetime import datetime, timedelta
import pandas as pd
from PIL import Image
import streamlit as st


import streamlit as st

# 1. Page Configuration
st.set_page_config(
    page_title="AI Study Notebook",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# 2. Custom CSS: Fixes Math LaTeX, Keeps Sidebar Buttons Inline, Locks Viewport
st.markdown(
    """
    <style>
    /* 1. LOCK VIEWPORT & PREVENT WOBBLE */
    html, body, [data-testid="stAppViewContainer"], .main {
        overflow-x: hidden !important;
        max-width: 100vw !important;
    }

    .block-container {
        padding-top: 3rem !important;
        padding-bottom: 5rem !important;
        padding-left: 0.5rem !important;
        padding-right: 0.5rem !important;
        max-width: 100% !important;
    }

    /* 2. FIX LATEX FORMULAS (Prevents vertical character stacking) */
    .katex, .katex *, .katex-display, .katex-display *, .MathJax, .MathJax * {
        white-space: nowrap !important;
        word-break: normal !important;
        overflow-wrap: normal !important;
        line-height: normal !important;
    }
    
    .katex-display {
        overflow-x: auto !important;
        overflow-y: hidden !important;
        padding: 0.4em 0 !important;
        margin: 0.5em 0 !important;
    }

    /* 3. FIX SIDEBAR THREADS (Forces Thread Name + Delete Button onto 1 Line) */
    [data-testid="stSidebar"] [data-testid="stHorizontalBlock"] {
        flex-direction: row !important;
        flex-wrap: nowrap !important;
        align-items: center !important;
        gap: 0.25rem !important;
    }

    [data-testid="stSidebar"] [data-testid="column"] {
        width: auto !important;
        min-width: 0 !important;
    }

    [data-testid="stSidebar"] [data-testid="column"]:first-child {
        flex: 1 1 80% !important;
    }

    [data-testid="stSidebar"] [data-testid="column"]:last-child {
        flex: 0 0 20% !important;
    }

    [data-testid="stSidebar"] .stButton > button {
        padding: 0.25rem 0.5rem !important;
        height: 38px !important;
        font-size: 14px !important;
        white-space: nowrap !important;
        text-overflow: ellipsis !important;
        overflow: hidden !important;
    }

    /* 4. SAFE TEXT WRAPPING FOR REGULAR PROSE ONLY */
    p, li, h1, h2, h3, h4 {
        white-space: pre-wrap !important;
        word-break: break-word !important;
        overflow-wrap: anywhere !important;
    }

    header[data-testid="stHeader"] {
        background-color: transparent !important;
        z-index: 99999 !important;
    }
    </style>
""",
    unsafe_allow_html=True,
)
# Optional sklearn dependency for local RAG vector similarity
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False

st.set_page_config(
    page_title="Agentic Study Platform - Job Prep Edition", page_icon="🎓", layout="wide"
)

DATA_FILE = "study_data.json"

# Updated list of verified active free models on OpenRouter
MODEL_OPTIONS = {
    "openrouter/free (Auto-Router - Highest Uptime)": "openrouter/free",
    "google/gemma-2-9b-it:free (Fast & Multimodal)": "google/gemma-2-9b-it:free",
    "meta-llama/llama-3.1-8b-instruct:free (High Accuracy)": (
        "meta-llama/llama-3.1-8b-instruct:free"
    ),
    "qwen/qwen-2.5-72b-instruct:free (Deep Reasoning & Math)": (
        "qwen/qwen-2.5-72b-instruct:free"
    ),
    "deepseek/deepseek-r1:free (Advanced Problem Solving)": (
        "deepseek/deepseek-r1:free"
    ),
    "nvidia/nemotron-3-ultra-550b-a55b:free (Complex Logic)": (
        "nvidia/nemotron-3-ultra-550b-a55b:free"
    ),
    "mistralai/mistral-7b-instruct:free (Fast Processing)": (
        "mistralai/mistral-7b-instruct:free"
    ),
}

# Fallback sequence for rate limits / 404 model unavailability
FALLBACK_MODELS = [
    "openrouter/free",
    "google/gemma-2-9b-it:free",
    "meta-llama/llama-3.1-8b-instruct:free",
    "qwen/qwen-2.5-72b-instruct:free",
    "mistralai/mistral-7b-instruct:free",
]

MATH_SYSTEM_PROMPT = {
    "role": "system",
    "content": (
        "You are an expert academic and job preparation assistant. When rendering formatting:\n"
        "1. ALWAYS place markdown headers (e.g., ### Section) on separate lines with BLANK LINES before and after.\n"
        "2. Put standalone display math equations inside double dollar signs on their OWN separate lines:\n"
        "$$\n"
        "f(x) = \\dots\n"
        "$$\n"
        "3. Wrap inline math variables/symbols in single dollar signs (e.g., $x = 5$).\n"
        "4. Render tables using valid Markdown table format with proper alignment rows (|---|---|).\n"
        "5. NEVER output orphan structural tags like \\end{cases} without a matching \\begin{cases}."
    ),
}


# --- ADVANCED LATEX & MARKDOWN FORMAT SANITIZER ---
def fix_latex_formatting(text: str) -> str:
    """Cleans raw LLM outputs, separates squished headers, and fixes KaTeX & Table rendering bugs."""
    if not text:
        return ""

    text = re.sub(r"(?<!\n)(###?\s+)", r"\n\n\1", text)
    text = re.sub(r"(\\end\{[a-zA-Z]+\})\s*(###?|---)", r"\1\n\n\2", text)
    text = re.sub(r"(?<![|\w\n])\n?^\s*---\s*$(?![|\w])", r"\n\n---\n\n", text, flags=re.MULTILINE)

    text = re.sub(r"\\\[\s*(.*?)\s*\\\]", r"\n$$\1$$\n", text, flags=re.DOTALL)
    text = re.sub(r"(?<!\w)\[\s*(\\.*?)\s*\]", r"\n$$\1$$\n", text, flags=re.DOTALL)
    text = re.sub(r"\\\(\s*(.*?)\s*\\\)", r"$\1$", text, flags=re.DOTALL)
    text = re.sub(r"(?<!\w)\(\s*(\\.*?)\s*\)", r"$\1$", text, flags=re.DOTALL)

    text = re.sub(
        r"\$\$\s*(\\begin\{(aligned|equation|gather|alignat|matrix|bmatrix|cases|array)\})",
        r"\n$$\n\1",
        text,
    )
    text = re.sub(
        r"(\\end\{(aligned|equation|gather|alignat|matrix|bmatrix|cases|array)\})\s*\$\$",
        r"\1\n$$\n",
        text,
    )

    if "\\end{cases}" in text and "\\begin{cases}" not in text:
        text = text.replace("\\end{cases}", "")

    return text


def clean_json_response(content: str) -> str:
    """Safely extracts raw JSON arrays or objects from markdown responses."""
    if not content:
        return ""
    cleaned = re.sub(r"^```(?:json)?", "", content.strip(), flags=re.MULTILINE)
    cleaned = re.sub(r"```$", "", cleaned.strip(), flags=re.MULTILINE)
    match = re.search(r"([\[\{].*[\]\}])", cleaned, re.DOTALL)
    return match.group(1) if match else cleaned.strip()


def get_openrouter_client(api_key):
    from openai import OpenAI

    return OpenAI(
        base_url="https://openrouter.ai/api/v1", api_key=api_key, timeout=45.0
    )


# --- RAG VECTOR RETRIEVAL ENGINE ---
def retrieve_relevant_chunks(documents: list, query: str, top_k: int = 5, chunk_size: int = 600) -> str:
    """Chunks text sources and retrieves top-k relevant segments using TF-IDF vector similarity."""
    all_chunks = []
    
    for doc in documents:
        text = doc.get("text", "")
        doc_name = doc.get("name", "Document")
        if not text or text.startswith("[Image Document"):
            continue
        
        for i in range(0, len(text), chunk_size - 100):
            chunk = text[i : i + chunk_size]
            if len(chunk.strip()) > 50:
                all_chunks.append({"source": doc_name, "chunk": chunk})

    if not all_chunks:
        return ""

    if len(all_chunks) <= top_k:
        return "\n\n".join([f"--- Context from {c['source']} ---\n{c['chunk']}" for c in all_chunks])

    chunk_texts = [c["chunk"] for c in all_chunks]

    if SKLEARN_AVAILABLE:
        try:
            vectorizer = TfidfVectorizer(stop_words="english")
            tfidf_matrix = vectorizer.fit_transform(chunk_texts + [query])
            cosine_sim = cosine_similarity(tfidf_matrix[-1:], tfidf_matrix[:-1]).flatten()
            top_indices = cosine_sim.argsort()[-top_k:][::-1]
            retrieved = [all_chunks[i] for i in top_indices if cosine_sim[i] > 0.05]
            if retrieved:
                return "\n\n".join([f"--- Relevant Snippet from {c['source']} ---\n{c['chunk']}" for c in retrieved])
        except Exception:
            pass

    query_words = set(query.lower().split())
    scored_chunks = []
    for c in all_chunks:
        score = sum(1 for w in query_words if w in c["chunk"].lower())
        scored_chunks.append((score, c))
    
    scored_chunks.sort(key=lambda x: x[0], reverse=True)
    top_chunks = [c[1] for c in scored_chunks[:top_k]]
    return "\n\n".join([f"--- Context from {c['source']} ---\n{c['chunk']}" for c in top_chunks])


# --- SUPERMEMO-2 (SM-2) SPACED REPETITION ALGORITHM ---
def calculate_sm2(quality: int, repetitions: int, interval: int, easiness_factor: float):
    if quality >= 3:
        if repetitions == 0:
            interval = 1
        elif repetitions == 1:
            interval = 6
        else:
            interval = math.ceil(interval * easiness_factor)
        repetitions += 1
    else:
        repetitions = 0
        interval = 1

    easiness_factor = easiness_factor + (0.1 - (5 - quality) * (0.08 + (5 - quality) * 0.02))
    if easiness_factor < 1.3:
        easiness_factor = 1.3

    next_date = (datetime.now() + timedelta(days=interval)).strftime("%Y-%m-%d")
    return repetitions, interval, round(easiness_factor, 2), next_date


# --- UNIVERSAL MODEL EXECUTION WITH FALLBACK ---
def execute_completion_with_fallback(client, messages, preferred_model, status_box=None):
    """Executes chat completion with automatic fallback retry loop across free models."""
    candidates = [preferred_model] + [m for m in FALLBACK_MODELS if m != preferred_model]
    last_err = None

    for model_candidate in candidates:
        try:
            if status_box:
                status_box.write(f"📡 Querying model: `{model_candidate}`...")
            
            res = client.chat.completions.create(
                model=model_candidate,
                messages=messages,
            )
            content = res.choices[0].message.content
            if content and content.strip():
                return content, model_candidate
        except Exception as ex:
            last_err = str(ex)
            if status_box:
                status_box.write(f"⚠️ `{model_candidate}` unavailable ({last_err[:60]}...). Retrying backup model...")
            continue

    raise Exception(f"All free models overloaded or unavailable. Last error: {last_err}")


# --- DATA PERSISTENCE ---
def load_data():
    default_structure = {
        "chats": {"Default Chat": []},
        "mcq_subjects": {
            "General Math & Logic": {
                "sources": [],
                "chat": [],
                "mistakes": [],
                "instructions": "",
                "flashcards": [],
            }
        },
        "written_subjects": {
            "Job Essay & Analytical": {
                "sources": [],
                "chat": [],
                "mistakes": [],
                "instructions": "Focus on structure, argument clarity, and precise data.",
            }
        },
        "viva_history": [],
        "analytics": [],
    }
    if os.path.exists(DATA_FILE):
        try:
            with open(DATA_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
                if not isinstance(data, dict):
                    return default_structure
                if "chats" not in data or not data["chats"]:
                    data["chats"] = {"Default Chat": []}
                if "mcq_subjects" not in data:
                    data["mcq_subjects"] = default_structure["mcq_subjects"]
                if "written_subjects" not in data:
                    data["written_subjects"] = default_structure["written_subjects"]
                if "viva_history" not in data:
                    data["viva_history"] = []
                if "analytics" not in data:
                    data["analytics"] = []
                return data
        except Exception:
            st.warning("⚠️ Could not read saved data. Initializing new state.")
    return default_structure


def save_data(data):
    try:
        with open(DATA_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.error(f"Failed to save data: {e}")


# --- STATE INITIALIZATION ---
if "db" not in st.session_state:
    st.session_state.db = load_data()

if "active_mode" not in st.session_state:
    st.session_state.active_mode = "general_chat"

if "active_chat" not in st.session_state:
    existing_chats = list(st.session_state.db["chats"].keys())
    st.session_state.active_chat = (
        existing_chats[0] if existing_chats else "Default Chat"
    )

if "trigger_regenerate" not in st.session_state:
    st.session_state.trigger_regenerate = False


def compress_image_to_b64(image_bytes, max_dim=1000, quality=75):
    try:
        img = Image.open(io.BytesIO(image_bytes))
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        img.thumbnail((max_dim, max_dim))
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=quality)
        b64_str = base64.b64encode(buf.getvalue()).decode("utf-8")
        return f"data:image/jpeg;base64,{b64_str}"
    except Exception:
        b64_str = base64.b64encode(image_bytes).decode("utf-8")
        return f"data:image/jpeg;base64,{b64_str}"


def extract_file_data(uploaded_file):
    file_type = uploaded_file.name.split(".")[-1].lower()
    extracted_text = ""
    image_url = None

    try:
        if file_type in ["png", "jpg", "jpeg"]:
            image_url = compress_image_to_b64(uploaded_file.getvalue())
            extracted_text = f"[Image Document: {uploaded_file.name}]"

        elif file_type == "pdf":
            pdf_text = []
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(uploaded_file.getvalue())) as pdf:
                    for page_idx, page in enumerate(pdf.pages):
                        if page_idx >= 50:
                            pdf_text.append("\n[Truncated remaining pages]")
                            break
                        page_content = f"--- Page {page_idx + 1} ---\n"
                        tables = page.extract_tables()
                        md_tables = ""
                        if tables:
                            for tbl in tables:
                                clean_rows = [[str(cell or '').strip().replace('\n', ' ') for cell in row] for row in tbl if any(row)]
                                if len(clean_rows) > 1:
                                    header = clean_rows[0]
                                    md_tables += "\n| " + " | ".join(header) + " |\n"
                                    md_tables += "| " + " | ".join(["---"] * len(header)) + " |\n"
                                    for row in clean_rows[1:]:
                                        row_padded = row + [""] * (len(header) - len(row))
                                        md_tables += "| " + " | ".join(row_padded[:len(header)]) + " |\n"
                                    md_tables += "\n"

                        raw_text = page.extract_text() or ""
                        if raw_text:
                            page_content += raw_text + "\n"
                        if md_tables:
                            page_content += "\n**Extracted Tables:**\n" + md_tables
                        if page_content.strip():
                            pdf_text.append(page_content)
            except Exception:
                from pypdf import PdfReader
                pdf_stream = io.BytesIO(uploaded_file.getvalue())
                reader = PdfReader(pdf_stream, strict=False)
                for page_idx, page in enumerate(reader.pages):
                    if page_idx >= 50:
                        pdf_text.append("\n[Truncated remaining pages]")
                        break
                    try:
                        txt = page.extract_text()
                        if txt:
                            pdf_text.append(f"--- Page {page_idx + 1} ---\n" + txt)
                    except Exception:
                        continue

            extracted_text = "\n\n".join(pdf_text) if pdf_text else "[Scanned/Unreadable PDF]"

        elif file_type == "docx":
            from docx import Document
            doc = Document(io.BytesIO(uploaded_file.getvalue()))
            extracted_text = "\n".join([p.text for p in doc.paragraphs])

        elif file_type in ["csv", "xlsx"]:
            df = pd.read_csv(uploaded_file) if file_type == "csv" else pd.read_excel(uploaded_file)
            extracted_text = df.head(100).to_string(index=False)
        else:
            extracted_text = uploaded_file.getvalue().decode("utf-8", errors="ignore")

    except Exception as e:
        extracted_text = f"Error extracting {uploaded_file.name}: {str(e)}"

    return {
        "name": uploaded_file.name,
        "type": file_type,
        "text": extracted_text,
        "image_url": image_url,
        "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }


def generate_docx(subject_name, mistakes_list, section_type="MCQ"):
    from docx import Document
    doc = Document()
    doc.add_heading(f"{section_type} Revision Guide: {subject_name}", level=0)
    if not mistakes_list:
        doc.add_paragraph("No weak spots logged yet.")
    for item in mistakes_list:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[{item.get('date', 'N/A')}] ").bold = True
        p.add_run(f"{item.get('concept', item.get('area', 'Topic'))}\n")
        p.add_run(f"Takeaway: {item.get('takeaway', item.get('correction', 'N/A'))}")
    filename = f"{subject_name}_{section_type}_Revision.docx"
    doc.save(filename)
    return filename


# ==========================================
# SIDEBAR NAVIGATION
# ==========================================
with st.sidebar:
    st.title("🎓 OpenRouter Studio")

    default_key = st.secrets.get("OPENROUTER_API_KEY", st.session_state.get("saved_openrouter_key", ""))
    api_key = st.text_input(
        "OpenRouter API Key",
        value=default_key,
        type="password",
        help="Get a free key from openrouter.ai",
    )

    if api_key:
        st.session_state.saved_openrouter_key = api_key
    else:
        st.warning("⚠️ Enter your free OpenRouter API Key to start.")

    st.divider()

    nav_choice = st.radio(
        "📌 Navigation",
        [
            "💬 General Chat",
            "📚 Notebook Workspaces",
            "🎙️ AI Mock Viva & Interview",
            "📊 Analytics & Mastery Dashboard",
        ],
    )
    
    mode_map = {
        "💬 General Chat": "general_chat",
        "📚 Notebook Workspaces": "notebook_studio",
        "🎙️ AI Mock Viva & Interview": "mock_viva",
        "📊 Analytics & Mastery Dashboard": "analytics_dash",
    }
    st.session_state.active_mode = mode_map[nav_choice]

    st.divider()

    if st.session_state.active_mode == "general_chat":
        col_c1, col_c2 = st.columns([3, 1])
        col_c1.subheader("Threads")

        if col_c2.button("➕ New", key="btn_new_chat"):
            new_title = f"Chat {datetime.now().strftime('%H:%M:%S')}"
            st.session_state.db["chats"][new_title] = []
            save_data(st.session_state.db)
            st.session_state.active_chat = new_title
            st.rerun()

        for chat_name in list(st.session_state.db["chats"].keys()):
            col_btn, col_del = st.columns([4, 1])
            is_active = chat_name == st.session_state.active_chat
            label = f"👉 {chat_name[:12]}" if is_active else f"💬 {chat_name[:12]}"

            if col_btn.button(label, key=f"sel_{chat_name}"):
                st.session_state.active_chat = chat_name
                st.rerun()

            if len(st.session_state.db["chats"]) > 1:
                if col_del.button("🗑️", key=f"del_{chat_name}"):
                    del st.session_state.db["chats"][chat_name]
                    st.session_state.active_chat = list(st.session_state.db["chats"].keys())[0]
                    save_data(st.session_state.db)
                    st.rerun()

        st.divider()
        rename_input = st.text_input("Rename Thread", st.session_state.active_chat)
        if st.button("Update Title") and rename_input and rename_input != st.session_state.active_chat:
            st.session_state.db["chats"][rename_input] = st.session_state.db["chats"].pop(st.session_state.active_chat)
            st.session_state.active_chat = rename_input
            save_data(st.session_state.db)
            st.rerun()


# ==========================================
# VIEW 1: GENERAL CHAT
# ==========================================
if st.session_state.active_mode == "general_chat":
    st.title(f"💬 {st.session_state.active_chat}")

    chat_history = st.session_state.db["chats"].get(st.session_state.active_chat, [])

    for idx, msg in enumerate(chat_history):
        with st.chat_message(msg["role"]):
            if msg.get("file_names"):
                st.caption(f"📎 Attached Context: {', '.join(msg['file_names'])}")
            st.markdown(fix_latex_formatting(msg["content"]))
            if msg["role"] == "assistant":
                with st.expander("📋 Copy Raw Text"):
                    st.code(msg["content"], language="markdown")

    if chat_history and chat_history[-1]["role"] in ["user", "assistant"]:
        col_reg, _ = st.columns([2, 5])
        if col_reg.button("🔄 Regenerate / Retry Response", key="btn_regen_chat"):
            st.session_state.trigger_regenerate = True
            st.rerun()

    st.markdown("---")

    col_ctrl1, col_ctrl2 = st.columns([1, 1])
    with col_ctrl1:
        selected_label = st.selectbox(
            "🌐 AI Model Architecture",
            options=list(MODEL_OPTIONS.keys()),
            index=0,
            key="chat_model_select",
        )
        selected_model_slug = MODEL_OPTIONS[selected_label]

    with col_ctrl2:
        attached_files = st.file_uploader(
            "Attach Context Sources (Multiple PDFs, DOCX, CSV, Images)",
            type=["png", "jpg", "jpeg", "pdf", "docx", "csv", "xlsx"],
            accept_multiple_files=True,
            key="gen_chat_files",
        )

    user_query = st.chat_input("Ask anything, request a math derivation, or submit a problem...")

    should_process = False
    if user_query:
        should_process = True
        st.session_state.trigger_regenerate = False
    elif st.session_state.trigger_regenerate and chat_history:
        should_process = True
        st.session_state.trigger_regenerate = False
        if chat_history[-1]["role"] == "assistant":
            chat_history.pop()
        if chat_history and chat_history[-1]["role"] == "user":
            user_query = chat_history[-1]["content"]

    if should_process and user_query:
        if not st.session_state.get("saved_openrouter_key"):
            st.error("Please enter an OpenRouter API key in the sidebar first!")
        else:
            client = get_openrouter_client(st.session_state.saved_openrouter_key)

            temp_file_text = ""
            image_urls = []
            file_names_list = []
            if attached_files:
                for f in attached_files:
                    f_data = extract_file_data(f)
                    file_names_list.append(f_data['name'])
                    temp_file_text += f"\n\n--- Attached Context: {f_data['name']} ---\n{f_data['text']}"
                    if f_data.get("image_url"):
                        image_urls.append(f_data["image_url"])

            if not (st.session_state.trigger_regenerate and chat_history[-1]["role"] == "user"):
                user_msg_record = {
                    "role": "user",
                    "content": user_query,
                    "context": temp_file_text,
                    "file_names": file_names_list
                }
                chat_history.append(user_msg_record)
                st.session_state.db["chats"][st.session_state.active_chat] = chat_history
                save_data(st.session_state.db)

            with st.chat_message("user"):
                if file_names_list:
                    st.caption(f"📎 Attached Context: {', '.join(file_names_list)}")
                st.markdown(user_query)

            with st.chat_message("assistant"):
                with st.status("🧠 Processing request...", expanded=True) as status:
                    response_placeholder = st.empty()
                    full_response = ""

                    api_messages = [MATH_SYSTEM_PROMPT]
                    for m in chat_history[-6:-1]:
                        api_messages.append({"role": m["role"], "content": m["content"]})

                    latest_prompt = f"{user_query}\n\n{temp_file_text}".strip()
                    if image_urls:
                        user_content = [{"type": "text", "text": latest_prompt}]
                        for img_url in image_urls[:2]:
                            user_content.append({"type": "image_url", "image_url": {"url": img_url}})
                        api_messages.append({"role": "user", "content": user_content})
                    else:
                        api_messages.append({"role": "user", "content": latest_prompt})

                    candidate_models = [selected_model_slug] + [
                        m for m in FALLBACK_MODELS if m != selected_model_slug
                    ]
                    
                    success = False
                    last_error = None

                    for model_candidate in candidate_models:
                        try:
                            status.write(f"📡 Querying model: `{model_candidate}`...")
                            response = client.chat.completions.create(
                                model=model_candidate, messages=api_messages, stream=True
                            )

                            for chunk in response:
                                if hasattr(chunk, "choices") and chunk.choices:
                                    choice = chunk.choices[0]
                                    if hasattr(choice, "delta") and choice.delta:
                                        content = getattr(choice.delta, "content", None)
                                        if content:
                                            full_response += content
                                            response_placeholder.markdown(
                                                fix_latex_formatting(full_response) + " ▌"
                                            )

                            cleaned_final = fix_latex_formatting(full_response)
                            if cleaned_final.strip():
                                response_placeholder.markdown(cleaned_final)
                                status.update(label=f"✅ Finished using `{model_candidate}`", state="complete")
                                chat_history.append({"role": "assistant", "content": cleaned_final})
                                st.session_state.db["chats"][st.session_state.active_chat] = chat_history
                                save_data(st.session_state.db)
                                success = True
                                break

                        except Exception as ex:
                            last_error = str(ex)
                            status.write(f"⚠️ `{model_candidate}` unavailable ({last_error[:60]}...). Trying backup model...")
                            full_response = ""

                    if not success:
                        status.update(label="❌ All Free Models Unavailable", state="error")
                        err_msg = f"⚠️ **Error Details**: `{last_error}`\n\nPlease retry in 30 seconds."
                        st.error(err_msg)
                        chat_history.append({"role": "assistant", "content": err_msg})
                        st.session_state.db["chats"][st.session_state.active_chat] = chat_history
                        save_data(st.session_state.db)

# ==========================================
# VIEW 2: NOTEBOOK WORKSPACES
# ==========================================
elif st.session_state.active_mode == "notebook_studio":
    st.title("📚 Persistent Notebook Workspaces")

    selected_label = st.selectbox(
        "🌐 Workspace Model (Vision & Document Capable)",
        options=list(MODEL_OPTIONS.keys()),
        index=0,
        key="notebook_model_select",
    )
    selected_model_slug = MODEL_OPTIONS[selected_label]

    workspace_type = st.radio(
        "Select Workspace Mode",
        ["📝 MCQ Workspace & Exam Simulator", "✍️ Focus Written Workspace"],
        horizontal=True,
    )

    # --------------------------------------------------------------------------
    # WORKSPACE 1: MCQ WORKSPACE & EXAM SIMULATOR
    # --------------------------------------------------------------------------
    if workspace_type == "📝 MCQ Workspace & Exam Simulator":
        with st.expander("⏱️ Practice & Exam Timer Settings", expanded=False):
            col_mcq_t1, col_mcq_t2 = st.columns([2, 1])
            mcq_mins = col_mcq_t1.number_input("Set Timer (Minutes)", 1, 180, 15, 5, key="mcq_timer_input")
            if col_mcq_t2.button("🚀 Start / Reset Timer", key="btn_mcq_timer"):
                st.session_state.mcq_timer_end = datetime.now().timestamp() + (mcq_mins * 60)
                st.success(f"Timer set for {mcq_mins} minutes!")

            if "mcq_timer_end" in st.session_state:
                remaining_sec = int(st.session_state.mcq_timer_end - datetime.now().timestamp())
                if remaining_sec > 0:
                    mins, secs = divmod(remaining_sec, 60)
                    st.info(f"⏳ **Time Remaining**: {mins:02d}m {secs:02d}s")
                else:
                    st.error("⏰ **Time is up! Submit your answers now.**")

        st.divider()

        col_s1, col_s2, col_s3 = st.columns([2, 1, 1])
        mcq_subs = list(st.session_state.db["mcq_subjects"].keys())
        selected_mcq_sub = col_s1.selectbox("Select Subject / Notebook", mcq_subs if mcq_subs else ["None"])

        new_mcq_sub = col_s2.text_input("New MCQ Notebook")
        if col_s2.button("Add Notebook", key="btn_add_mcq_sub") and new_mcq_sub:
            st.session_state.db["mcq_subjects"][new_mcq_sub] = {
                "sources": [],
                "chat": [],
                "mistakes": [],
                "instructions": "",
                "flashcards": [],
            }
            save_data(st.session_state.db)
            st.rerun()

        if selected_mcq_sub and selected_mcq_sub != "None":
            if col_s3.button("🗑️ Delete Notebook", key="btn_del_mcq_sub"):
                del st.session_state.db["mcq_subjects"][selected_mcq_sub]
                save_data(st.session_state.db)
                st.success(f"Deleted notebook '{selected_mcq_sub}'!")
                st.rerun()

            sub_data = st.session_state.db["mcq_subjects"][selected_mcq_sub]
            if "sources" not in sub_data:
                sub_data["sources"] = []
            if "flashcards" not in sub_data:
                sub_data["flashcards"] = []

            with st.expander(f"📚 Persistent Sources for '{selected_mcq_sub}' ({len(sub_data['sources'])} Saved)", expanded=True):
                st.markdown("Upload files here once. Local Vector Retrieval automatically chunks and retrieves context.")
                new_mcq_files = st.file_uploader(
                    "Add Sources to Notebook (PDF, DOCX, CSV, Images)",
                    type=["pdf", "docx", "csv", "xlsx", "png", "jpg", "jpeg"],
                    accept_multiple_files=True,
                    key=f"mcq_sources_up_{selected_mcq_sub}",
                )

                if st.button("💾 Save Files to Notebook Sources", key=f"btn_save_mcq_sources_{selected_mcq_sub}"):
                    if new_mcq_files:
                        for nf in new_mcq_files:
                            extracted_obj = extract_file_data(nf)
                            sub_data["sources"].append(extracted_obj)
                        save_data(st.session_state.db)
                        st.success("Saved sources to notebook!")
                        st.rerun()

                if sub_data["sources"]:
                    st.write("---")
                    st.write("**Current Saved Sources in Notebook:**")
                    for s_idx, src in enumerate(sub_data["sources"]):
                        sc1, sc2 = st.columns([5, 1])
                        sc1.markdown(f"📄 **{src['name']}** *(Added {src['date']})*")
                        if sc2.button("🗑️ Delete", key=f"del_mcq_src_{selected_mcq_sub}_{s_idx}"):
                            sub_data["sources"].pop(s_idx)
                            save_data(st.session_state.db)
                            st.rerun()

            m_tab1, m_tab2, m_tab3 = st.tabs([
                "🎯 Quiz Generator & Exam Simulator", 
                "📖 Mistakes Register & Notes", 
                "🎴 Spaced Repetition Flashcards (Anki SM-2)"
            ])

            with m_tab1:
                col_cfg1, col_cfg2, col_cfg3, col_cfg4 = st.columns(4)
                num_q = col_cfg1.number_input("Number of Questions", 1, 30, 5)
                diff = col_cfg2.selectbox("Difficulty", ["Medium", "Hard", "Advanced Exam"])
                lang_choice = col_cfg3.selectbox("Language", ["English", "Bangla (বাংলা)", "Bilingual"])
                
                neg_marking_enabled = col_cfg4.checkbox("Negative Marking Exam Mode", value=True)
                penalty_rate = 0.25 if neg_marking_enabled else 0.0

                custom_instructions = st.text_area(
                    "Chapter / Custom Prompt Instructions for MCQ Generation",
                    value=sub_data.get("instructions", "Focus on core chapter concepts and past weak points."),
                    height=80,
                    key="mcq_inst_input"
                )
                if st.button("💾 Save Instructions"):
                    sub_data["instructions"] = custom_instructions
                    save_data(st.session_state.db)
                    st.success("Instructions saved!")

                if st.button("🚀 Generate Chapter Quiz (with RAG Context)"):
                    if not st.session_state.get("saved_openrouter_key"):
                        st.error("Please enter your API Key in the sidebar.")
                    else:
                        client = get_openrouter_client(st.session_state.saved_openrouter_key)

                        rag_context = retrieve_relevant_chunks(
                            sub_data.get("sources", []), 
                            query=f"{selected_mcq_sub} {custom_instructions}",
                            top_k=6
                        )

                        prior_mistakes = [
                            f"- {m.get('concept', '')}: {m.get('takeaway', '')}"
                            for m in sub_data.get("mistakes", [])
                        ]
                        weakness_context = "\n".join(prior_mistakes) if prior_mistakes else "None recorded yet."

                        prompt_text = f"""Language Requirement: {lang_choice}
Subject Notebook: {selected_mcq_sub}
Difficulty Level: {diff}
Target Questions Count: {num_q}

CUSTOM INSTRUCTIONS / CHAPTER FOCUS:
{custom_instructions}

LOGGED PAST MISTAKES & WEAKNESSES:
{weakness_context}

RETRIEVED VECTOR CONTEXT / SOURCE MATERIAL:
{rag_context if rag_context else "No vector sources. Generate standard exam questions for subject domain."}

GENERATE A QUIZ FOLLOWING THESE RULES:
1. Target chapter concepts in saved vector materials and past logged weaknesses.
2. DO NOT repeat exact duplicate questions from past mistakes; create NEW variations testing those concepts.
3. Use proper LaTeX notation ($...$ inline, $$...$$ block) for math and valid Markdown tables.
4. Output STRICT raw JSON array format without markdown backticks:
[
  {{"id": 1, "question": "...", "options": ["A) ...", "B) ...", "C) ...", "D) ..."], "correct": "A) ...", "explanation": "..."}}
]
"""

                        with st.status("🧠 Retrieving Context & Generating Quiz...", expanded=True) as sbox:
                            try:
                                response_text, used_model = execute_completion_with_fallback(
                                    client,
                                    [MATH_SYSTEM_PROMPT, {"role": "user", "content": prompt_text}],
                                    selected_model_slug,
                                    status_box=sbox
                                )
                                raw_json = clean_json_response(response_text)
                                st.session_state.quiz = json.loads(raw_json)
                                sbox.update(label=f"✅ Quiz Generated via `{used_model}`", state="complete")
                            except Exception as qe:
                                st.error(f"Quiz Generation Error: {str(qe)}")

                if "quiz" in st.session_state:
                    st.divider()
                    user_ans = {}
                    with st.form("mcq_form"):
                        for q in st.session_state.quiz:
                            st.write(f"**Q{q['id']}: {fix_latex_formatting(q['question'])}**")
                            opts = ["Skipped"] + q["options"]
                            user_ans[q["id"]] = st.radio(
                                f"Select Option for Q{q['id']}", opts, key=f"q_{q['id']}"
                            )
                            st.markdown("---")
                        submit_m = st.form_submit_button("Submit Exam Answers")

                    if submit_m:
                        correct_cnt = 0
                        wrong_cnt = 0
                        skipped_cnt = 0
                        existing_concepts = [m["concept"] for m in sub_data.get("mistakes", [])]

                        for q in st.session_state.quiz:
                            selected = user_ans[q["id"]]
                            if selected == "Skipped":
                                skipped_cnt += 1
                                st.warning(f"Q{q['id']} Skipped. Correct answer: {q['correct']}")
                            elif selected == q["correct"]:
                                correct_cnt += 1
                                st.success(f"Q{q['id']} Correct! 🎉")
                            else:
                                wrong_cnt += 1
                                st.error(f"Q{q['id']} Incorrect (Selected: {selected}). Correct: {q['correct']}")
                                st.info(f"💡 Explanation: {fix_latex_formatting(q['explanation'])}")

                                if q["question"] not in existing_concepts:
                                    sub_data["mistakes"].append({
                                        "date": datetime.now().strftime("%Y-%m-%d"),
                                        "concept": q["question"],
                                        "takeaway": q["explanation"],
                                    })
                                    sub_data["flashcards"].append({
                                        "id": len(sub_data["flashcards"]) + 1,
                                        "front": q["question"],
                                        "back": f"Correct Option: {q['correct']}\n\n{q['explanation']}",
                                        "repetitions": 0,
                                        "interval": 1,
                                        "easiness_factor": 2.5,
                                        "next_review": datetime.now().strftime("%Y-%m-%d"),
                                    })

                        net_score = correct_cnt - (wrong_cnt * penalty_rate)
                        total_q = len(st.session_state.quiz)
                        accuracy = (correct_cnt / (correct_cnt + wrong_cnt) * 100) if (correct_cnt + wrong_cnt) > 0 else 0

                        st.markdown("### 📊 Exam Performance Summary")
                        mcol1, mcol2, mcol3, mcol4 = st.columns(4)
                        mcol1.metric("Net Marks", f"{net_score:.2f} / {total_q}")
                        mcol2.metric("Accuracy", f"{accuracy:.1f}%")
                        mcol3.metric("Correct / Wrong", f"{correct_cnt} / {wrong_cnt}")
                        mcol4.metric("Penalty Deduction", f"-{wrong_cnt * penalty_rate:.2f}")

                        st.session_state.db["analytics"].append({
                            "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
                            "subject": selected_mcq_sub,
                            "type": "MCQ Exam",
                            "score": net_score,
                            "total": total_q,
                            "accuracy": round(accuracy, 1),
                        })
                        save_data(st.session_state.db)

            with m_tab2:
                st.subheader(f"Logged Mistakes Register ({selected_mcq_sub})")

                with st.expander("➕ Manually Log Weak Spot / Concept"):
                    manual_concept = st.text_input("Concept / Question")
                    manual_takeaway = st.text_area("Correct Explanation / Key Takeaway")
                    if st.button("Save Weak Spot"):
                        if manual_concept:
                            sub_data["mistakes"].append({
                                "date": datetime.now().strftime("%Y-%m-%d"),
                                "concept": manual_concept,
                                "takeaway": manual_takeaway
                            })
                            sub_data["flashcards"].append({
                                "id": len(sub_data["flashcards"]) + 1,
                                "front": manual_concept,
                                "back": manual_takeaway,
                                "repetitions": 0,
                                "interval": 1,
                                "easiness_factor": 2.5,
                                "next_review": datetime.now().strftime("%Y-%m-%d"),
                            })
                            save_data(st.session_state.db)
                            st.success("Logged into revision register and flashcard deck!")
                            st.rerun()

                st.markdown("---")
                for idx, m in enumerate(sub_data.get("mistakes", [])):
                    with st.expander(f"#{idx+1} [{m.get('date', 'N/A')}] {m['concept'][:60]}..."):
                        st.markdown(f"**Question / Concept:**\n{fix_latex_formatting(m['concept'])}")
                        st.markdown(f"**Explanation / Revision Takeaway:**\n{fix_latex_formatting(m['takeaway'])}")

                st.divider()
                if st.button("Export MCQ Revision Guide (.docx)"):
                    fpath = generate_docx(selected_mcq_sub, sub_data["mistakes"], "MCQ")
                    with open(fpath, "rb") as fp:
                        st.download_button("📥 Download Word Revision Guide", fp, file_name=fpath)

            with m_tab3:
                st.subheader("🎴 Spaced Repetition Flashcard Deck (Anki SM-2 Algorithm)")
                today_str = datetime.now().strftime("%Y-%m-%d")
                due_cards = [
                    card for card in sub_data.get("flashcards", [])
                    if card.get("next_review", today_str) <= today_str
                ]

                st.info(f"📌 **Cards Due for Review Today**: {len(due_cards)} / Total Deck: {len(sub_data.get('flashcards', []))}")

                if due_cards:
                    card = due_cards[0]
                    with st.container(border=True):
                        st.markdown(f"### **Question / Front:**\n{fix_latex_formatting(card['front'])}")
                        
                        if st.button("👁️ Reveal Answer / Back", key=f"reveal_{card['id']}"):
                            st.session_state[f"show_ans_{card['id']}"] = True

                        if st.session_state.get(f"show_ans_{card['id']}", False):
                            st.markdown("---")
                            st.markdown(f"### **Answer / Back:**\n{fix_latex_formatting(card['back'])}")
                            
                            st.markdown("**Rate Recall Difficulty (SM-2):**")
                            fcol0, fcol1, fcol2, fcol3 = st.columns(4)
                            
                            if fcol0.button("❌ Again (0)", key=f"sm_0_{card['id']}"):
                                card["repetitions"], card["interval"], card["easiness_factor"], card["next_review"] = calculate_sm2(
                                    0, card["repetitions"], card["interval"], card["easiness_factor"]
                                )
                                save_data(st.session_state.db)
                                st.rerun()

                            if fcol1.button("⚠️ Hard (3)", key=f"sm_3_{card['id']}"):
                                card["repetitions"], card["interval"], card["easiness_factor"], card["next_review"] = calculate_sm2(
                                    3, card["repetitions"], card["interval"], card["easiness_factor"]
                                )
                                save_data(st.session_state.db)
                                st.rerun()

                            if fcol2.button("👍 Good (4)", key=f"sm_4_{card['id']}"):
                                card["repetitions"], card["interval"], card["easiness_factor"], card["next_review"] = calculate_sm2(
                                    4, card["repetitions"], card["interval"], card["easiness_factor"]
                                )
                                save_data(st.session_state.db)
                                st.rerun()

                            if fcol3.button("🌟 Easy (5)", key=f"sm_5_{card['id']}"):
                                card["repetitions"], card["interval"], card["easiness_factor"], card["next_review"] = calculate_sm2(
                                    5, card["repetitions"], card["interval"], card["easiness_factor"]
                                )
                                save_data(st.session_state.db)
                                st.rerun()
                else:
                    st.success("🎉 All flashcards for this subject are reviewed for today!")

    # --------------------------------------------------------------------------
    # WORKSPACE 2: FOCUS WRITTEN WORKSPACE
    # --------------------------------------------------------------------------
    elif workspace_type == "✍️ Focus Written Workspace":
        with st.expander("⏱️ Written Exam Timer", expanded=False):
            col_w_t1, col_w_t2 = st.columns([2, 1])
            written_mins = col_w_t1.number_input("Set Timer (Minutes)", 1, 180, 30, 5, key="written_timer_input")
            if col_w_t2.button("🚀 Start / Reset Timer", key="btn_written_timer"):
                st.session_state.written_timer_end = datetime.now().timestamp() + (written_mins * 60)
                st.success(f"Timer set for {written_mins} minutes!")

            if "written_timer_end" in st.session_state:
                remaining_sec = int(st.session_state.written_timer_end - datetime.now().timestamp())
                if remaining_sec > 0:
                    mins, secs = divmod(remaining_sec, 60)
                    st.info(f"⏳ **Time Remaining**: {mins:02d}m {secs:02d}s")
                else:
                    st.error("⏰ **Time is up! Submit your answer now.**")

        st.divider()

        col_w1, col_w2, col_w3 = st.columns([2, 1, 1])
        w_subs = list(st.session_state.db["written_subjects"].keys())
        selected_w_sub = col_w1.selectbox("Select Subject / Notebook", w_subs if w_subs else ["None"])

        new_w_sub = col_w2.text_input("New Written Notebook")
        if col_w2.button("Add Notebook", key="btn_add_w_sub") and new_w_sub:
            st.session_state.db["written_subjects"][new_w_sub] = {
                "sources": [],
                "chat": [],
                "mistakes": [],
                "instructions": "Focus on clarity, analytical depth, and structural coherence.",
            }
            save_data(st.session_state.db)
            st.rerun()

        if selected_w_sub and selected_w_sub != "None":
            if col_w3.button("🗑️ Delete Notebook", key="btn_del_w_sub"):
                del st.session_state.db["written_subjects"][selected_w_sub]
                save_data(st.session_state.db)
                st.success(f"Deleted notebook '{selected_w_sub}'!")
                st.rerun()

            w_sub_data = st.session_state.db["written_subjects"][selected_w_sub]
            if "sources" not in w_sub_data:
                w_sub_data["sources"] = []

            with st.expander(f"📚 Persistent Sources for '{selected_w_sub}' ({len(w_sub_data['sources'])} Saved)", expanded=True):
                st.markdown("Upload reference documents or topic guidelines here.")
                new_w_files = st.file_uploader(
                    "Add Sources to Notebook (PDF, DOCX, Images)",
                    type=["pdf", "docx", "png", "jpg", "jpeg"],
                    accept_multiple_files=True,
                    key=f"written_sources_up_{selected_w_sub}",
                )

                if st.button("💾 Save Files to Notebook Sources", key=f"btn_save_w_sources_{selected_w_sub}"):
                    if new_w_files:
                        for nf in new_w_files:
                            extracted_obj = extract_file_data(nf)
                            w_sub_data["sources"].append(extracted_obj)
                        save_data(st.session_state.db)
                        st.success("Saved sources to notebook!")
                        st.rerun()

                if w_sub_data["sources"]:
                    st.write("---")
                    st.write("**Current Saved Sources in Notebook:**")
                    for s_idx, src in enumerate(w_sub_data["sources"]):
                        sc1, sc2 = st.columns([5, 1])
                        sc1.markdown(f"📄 **{src['name']}** *(Added {src['date']})*")
                        if sc2.button("🗑️ Delete", key=f"del_w_src_{selected_w_sub}_{s_idx}"):
                            w_sub_data["sources"].pop(s_idx)
                            save_data(st.session_state.db)
                            st.rerun()

            col_lang1, col_lang2 = st.columns([1, 2])
            lang_w_choice = col_lang1.selectbox("Evaluation Language", ["English", "Bangla (বাংলা)", "Bilingual"])

            custom_instructions = st.text_area(
                "Evaluation Rubric / Criteria Instructions",
                value=w_sub_data.get("instructions", "Focus on clarity, analytical depth, logic, and formula accuracy."),
                height=80,
            )
            if st.button("💾 Save Rubric Criteria"):
                w_sub_data["instructions"] = custom_instructions
                save_data(st.session_state.db)
                st.success("Rubric saved!")

            essay_input = st.text_area("Write or Paste Solution / Essay Submission", height=220)

            if st.button("🔍 Evaluate Written Solution"):
                if not st.session_state.get("saved_openrouter_key"):
                    st.error("Please enter your API Key in the sidebar.")
                elif not essay_input.strip() and not w_sub_data["sources"]:
                    st.warning("Please enter a written solution or ensure you have saved notebook sources.")
                else:
                    client = get_openrouter_client(st.session_state.saved_openrouter_key)

                    rag_context = retrieve_relevant_chunks(
                        w_sub_data.get("sources", []),
                        query=essay_input[:300],
                        top_k=5
                    )

                    prior_mistakes = [
                        f"- {m.get('area', '')}: {m.get('correction', '')}"
                        for m in w_sub_data.get("mistakes", [])
                    ]
                    weakness_context = "\n".join(prior_mistakes) if prior_mistakes else "None recorded yet."

                    prompt_text = f"""Language Requirement: {lang_w_choice}
Subject Notebook: {selected_w_sub}
Evaluation Criteria: {custom_instructions}

PAST LOGGED MISTAKES FOR THIS STUDENT:
{weakness_context}

RETRIEVED REFERENCE SOURCES:
{rag_context}

STUDENT WRITTEN SUBMISSION:
{essay_input}

Provide detailed feedback, point out errors/mistakes, and output STRICT JSON format:
{{
  "score": "85%",
  "weakness": "Key area needing improvement...",
  "strategy": "Actionable takeaway to fix mistake..."
}}
"""

                    with st.status("🧠 Evaluating Solution against Rubric & Vector Context...", expanded=True) as sbox:
                        try:
                            response_text, used_model = execute_completion_with_fallback(
                                client,
                                [MATH_SYSTEM_PROMPT, {"role": "user", "content": prompt_text}],
                                selected_model_slug,
                                status_box=sbox
                            )
                            raw_json = clean_json_response(response_text)
                            eval_data = json.loads(raw_json)

                            st.metric("Evaluation Score", eval_data["score"])
                            st.info(f"**Identified Weakness**: {fix_latex_formatting(eval_data['weakness'])}")
                            st.success(f"**Improvement Strategy**: {fix_latex_formatting(eval_data['strategy'])}")

                            w_sub_data["mistakes"].append({
                                "date": datetime.now().strftime("%Y-%m-%d"),
                                "area": eval_data["weakness"],
                                "correction": eval_data["strategy"],
                            })
                            st.session_state.db["analytics"].append({
                                "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
                                "subject": selected_w_sub,
                                "type": "Written Essay",
                                "score": eval_data["score"],
                                "total": "100%",
                                "accuracy": eval_data["score"],
                            })
                            save_data(st.session_state.db)
                            sbox.update(label=f"✅ Evaluation Complete via `{used_model}`", state="complete")
                        except Exception as we:
                            st.error(f"Evaluation Error: {str(we)}")

            st.divider()
            st.subheader(f"Logged Written Feedback Register ({selected_w_sub})")
            for idx, item in enumerate(w_sub_data.get("mistakes", [])):
                with st.expander(f"#{idx+1} [{item.get('date', 'N/A')}] {item.get('area', 'Feedback')[:60]}..."):
                    st.markdown(f"**Weakness Area:**\n{fix_latex_formatting(item.get('area', ''))}")
                    st.markdown(f"**Correction Strategy:**\n{fix_latex_formatting(item.get('correction', ''))}")

            st.divider()
            if st.button("Export Written Revision Guide (.docx)"):
                fpath = generate_docx(selected_w_sub, w_sub_data["mistakes"], "Written")
                with open(fpath, "rb") as fp:
                    st.download_button("📥 Download Revision Guide", fp, file_name=fpath)


# ==========================================
# VIEW 3: AI MOCK VIVA & INTERVIEW SIMULATOR
# ==========================================
elif st.session_state.active_mode == "mock_viva":
    st.title("🎙️ AI Mock Viva & Interview Simulator")
    st.markdown("Practice job oral interviews, BCS preliminary viva, bank candidate evaluations, or engineering technical boards.")

    col_v1, col_v2 = st.columns([1, 1])
    with col_v1:
        viva_role = st.selectbox(
            "Target Job Board / Role",
            [
                "BCS Administrative & General Cadre Viva",
                "Bank Officer & Financial Analyst Oral Exam",
                "Software Engineer & System Design Interview",
                "General Knowledge & Current Affairs Board",
                "Custom Job Domain",
            ],
        )
    with col_v2:
        viva_model_label = st.selectbox(
            "Interviewer Model",
            options=list(MODEL_OPTIONS.keys()),
            index=0,
            key="viva_model_select",
        )
        viva_model_slug = MODEL_OPTIONS[viva_model_label]

    if "viva_messages" not in st.session_state:
        st.session_state.viva_messages = []

    if st.button("🚀 Start New Viva Session"):
        st.session_state.viva_messages = [
            {
                "role": "assistant",
                "content": f"Welcome to the **{viva_role}** simulation. I am your board chairman today. Please introduce yourself and brief us on your academic background.",
            }
        ]
        st.rerun()

    st.divider()

    for v_msg in st.session_state.viva_messages:
        with st.chat_message(v_msg["role"]):
            st.markdown(fix_latex_formatting(v_msg["content"]))

    candidate_reply = st.chat_input("Speak / Type your candidate response here...")

    if candidate_reply:
        if not st.session_state.get("saved_openrouter_key"):
            st.error("Please enter your API Key in the sidebar.")
        else:
            client = get_openrouter_client(st.session_state.saved_openrouter_key)
            st.session_state.viva_messages.append({"role": "user", "content": candidate_reply})

            with st.chat_message("user"):
                st.markdown(candidate_reply)

            with st.chat_message("assistant"):
                with st.status("🧠 Viva Chairman is evaluating response...", expanded=True) as sbox:
                    viva_prompt = f"""You are a strict, formal Viva Board Examiner for: {viva_role}.
Evaluate the candidate's last answer, point out conciseness, factual accuracy, or tone flaws, give a quick mark out of 10, and ask the NEXT follow-up question.

Maintain strict board atmosphere."""

                    messages_payload = [{"role": "system", "content": viva_prompt}] + st.session_state.viva_messages[-6:]

                    try:
                        reply_content, used_model = execute_completion_with_fallback(
                            client,
                            messages_payload,
                            viva_model_slug,
                            status_box=sbox
                        )
                        st.markdown(fix_latex_formatting(reply_content))
                        st.session_state.viva_messages.append({"role": "assistant", "content": reply_content})
                        sbox.update(label=f"✅ Response evaluated via `{used_model}`", state="complete")
                    except Exception as ve:
                        st.error(f"Viva Error: {str(ve)}")


# ==========================================
# VIEW 4: ANALYTICS & MASTERY DASHBOARD
# ==========================================
elif st.session_state.active_mode == "analytics_dash":
    st.title("📊 Job Prep Analytics & Mastery Dashboard")
    st.markdown("Real-time performance tracking across all subjects, exams, and flashcards.")

    history_records = st.session_state.db.get("analytics", [])

    total_exams = len(history_records)
    total_mistakes = sum(
        len(sub.get("mistakes", [])) for sub in st.session_state.db["mcq_subjects"].values()
    ) + sum(
        len(sub.get("mistakes", [])) for sub in st.session_state.db["written_subjects"].values()
    )
    
    total_flashcards = sum(
        len(sub.get("flashcards", [])) for sub in st.session_state.db["mcq_subjects"].values()
    )

    col_a1, col_a2, col_a3, col_a4 = st.columns(4)
    col_a1.metric("Total Quizzes / Exams Completed", total_exams)
    col_a2.metric("Total Logged Weaknesses", total_mistakes)
    col_a3.metric("Spaced Repetition Deck Size", total_flashcards)
    col_a4.metric("Saved Notebook Workspaces", len(st.session_state.db["mcq_subjects"]) + len(st.session_state.db["written_subjects"]))

    st.divider()

    if history_records:
        st.subheader("📈 Performance Trend History")
        df = pd.DataFrame(history_records)
        st.dataframe(df, use_container_width=True)

        if "accuracy" in df.columns:
            st.subheader("🎯 Accuracy Progression Over Time")
            st.line_chart(df, x="date", y="accuracy")
    else:
        st.info("💡 No exam history recorded yet. Complete quizzes or written submissions in Notebook Workspaces to populate analytics.")
